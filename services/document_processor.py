import os
import uuid
import PyPDF2
from google.cloud import vision
from typing import List, Dict, Any
from PIL import Image
import docx
import pandas as pd
import io
import json

from models import FileType, ItemDetail
from services.ai_service import RFQGenerator
from services.read_pdf import AzureAIPDFReader
from config import settings

os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r'C:\Users\naman\Desktop\up\test\ocrrfq-4d49980049c0.json'

# Azure Document Intelligence credentials from config
AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT = settings.AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT
AZURE_DOCUMENT_INTELLIGENCE_KEY = settings.AZURE_DOCUMENT_INTELLIGENCE_KEY

def analyze_text_patterns(text: str) -> None:
    """Analyze and print basic patterns found in the extracted text"""
    if not text or len(text) < 10:
        print("Text too short for pattern analysis")
        return
    
    # Check for likely patterns
    patterns = []
    if any(word in text.lower() for word in ["qty", "quantity", "pcs", "units"]):
        patterns.append("Quantity indicators")
    
    if any(word in text.lower() for word in ["$", "usd", "eur", "price", "cost", "rate"]):
        patterns.append("Price/cost indicators")
    
    if any(word in text.lower() for word in ["spec", "specification", "dimension", "size", "weight"]):
        patterns.append("Specifications")
    
    if any(word in text.lower() for word in ["model", "part", "no.", "number", "sku", "code"]):
        patterns.append("Part/model numbers")
        
    if any(word in text.lower() for word in ["brand", "manufacturer", "vendor", "supplier"]):
        patterns.append("Brand/vendor references")
    
    # Print found patterns
    if patterns:
        print("Detected patterns in the text:")
        for pattern in patterns:
            print(f"  - {pattern}")
    else:
        print("No specific procurement patterns detected in the text")

async def process_document(file_path: str, file_type: FileType) -> List[ItemDetail]:
    """
    Process an uploaded document, extract text content, and use AI to generate RFQ items
    
    Args:
        file_path: Path to the uploaded file
        file_type: Type of the file (PDF, DOCX, EXCEL, IMAGE)
        
    Returns:
        List of extracted items as identified by the AI service
    """
    extracted_text = ""
    
    try:
        # Extract text based on file type
        if file_type == FileType.PDF:
            extracted_text = extract_text_from_pdf(file_path)
            print(f"\n--- EXTRACTED PDF CONTENT FROM {os.path.basename(file_path)} ---")
            print(f"First 500 chars: {extracted_text[:500]}...")
            print(f"Total length: {len(extracted_text)} characters")
        elif file_type == FileType.DOCX:
            extracted_text = extract_text_from_docx(file_path)
            print(f"\n--- EXTRACTED DOCX CONTENT FROM {os.path.basename(file_path)} ---")
            print(f"First 500 chars: {extracted_text[:500]}...")
            print(f"Total length: {len(extracted_text)} characters")
        elif file_type == FileType.EXCEL:
            extracted_text = extract_text_from_excel(file_path)
            print(f"\n--- EXTRACTED EXCEL CONTENT FROM {os.path.basename(file_path)} ---")
            print(f"First 500 chars: {extracted_text[:500]}...")
            print(f"Total length: {len(extracted_text)} characters")
        elif file_type == FileType.IMAGE:
            # For images, use OCR to extract text
            extracted_text = extract_text_from_image(file_path)
            print(f"\n--- EXTRACTED IMAGE CONTENT FROM {os.path.basename(file_path)} ---")
            print(f"OCR text: {extracted_text}")
            print(f"Total length: {len(extracted_text)} characters")
        
        # Analyze text patterns
        if extracted_text:
            analyze_text_patterns(extracted_text)
            
            # Generate RFQ using AI service
            print("\n--- GENERATING RFQ ITEMS WITH AI ---")
            rfq_generator = RFQGenerator()
            ai_result = rfq_generator.generate_rfq(extracted_text)
            
            try:
                # Parse the AI response 
                result_data = json.loads(ai_result)
                items = []
                
                print(f"AI identified {len(result_data.get('items', []))} items in the document")
                
                # Convert AI results to ItemDetail objects
                for item_data in result_data.get('items', []):
                    item = ItemDetail(
                        id=str(uuid.uuid4()),
                        name=item_data.get('name', f"Item from {os.path.basename(file_path)}"),
                        quantity=item_data.get('quantity', 1),
                        description=item_data.get('description', '')
                    )
                    items.append(item)
                    print(f"- Item: {item.name}, Quantity: {item.quantity}")
                
                if items:
                    return items
            except json.JSONDecodeError as e:
                print(f"Error parsing AI response: {e}")
                print(f"Raw AI response: {ai_result}")
        
        # Fallback: if AI processing fails, return a single item with raw content
        if extracted_text:
            item = ItemDetail(
                id=str(uuid.uuid4()),
                name=f"Document Content: {os.path.basename(file_path)}",
                description=extracted_text
            )
            print("Falling back to raw content as a single item")
            return [item]
    except Exception as e:
        print(f"Error processing document: {e}")
    
    # Return empty list if processing failed
    return []

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using Azure AI Document Intelligence"""
    text = ""
    try:
        # Check if Azure credentials are available
        if AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT and AZURE_DOCUMENT_INTELLIGENCE_KEY:
            # Use Azure AI PDF Reader
            pdf_reader = AzureAIPDFReader(
                endpoint=AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT, 
                key=AZURE_DOCUMENT_INTELLIGENCE_KEY
            )
            
            # Extract structured content from PDF
            result = pdf_reader.extract_text(file_path)
            
            # Extract paragraphs and organize them into text
            if result and "paragraphs" in result:
                paragraphs = [p["text"] for p in result["paragraphs"]]
                text = "\n\n".join(paragraphs)
            
            # If no paragraphs, try to get text from page lines
            if not text and "pages" in result:
                for page in result["pages"]:
                    if "lines" in page:
                        text += "\n".join(page["lines"]) + "\n\n"
            
            # Include table data
            if "tables" in result:
                text += "\n\nTABLES:\n"
                for table in result["tables"]:
                    if "grid" in table:
                        for row in table["grid"]:
                            text += " | ".join([cell if cell else "" for cell in row]) + "\n"
                        text += "\n"
            
            print("Used Azure AI Document Intelligence for PDF extraction")
        else:
            # Fallback to PyPDF2 if Azure credentials are not available
            print("Azure credentials not found. Using PyPDF2 fallback for PDF extraction")
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        # Fallback to PyPDF2 if Azure extraction fails
        try:
            print("Falling back to PyPDF2 for PDF extraction")
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
        except Exception as inner_e:
            print(f"Fallback extraction also failed: {inner_e}")
    
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file"""
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
    return text

def extract_text_from_excel(file_path: str) -> str:
    """Extract text from an Excel file"""
    text = ""
    try:
        df = pd.read_excel(file_path)
        buffer = io.StringIO()
        df.to_csv(buffer)
        text = buffer.getvalue()
    except Exception as e:
        print(f"Error extracting text from Excel: {e}")
    return text

def extract_text_from_image(file_path: str) -> str:
    """Extract text from an image file using OCR"""
    text = ""
    client = vision.ImageAnnotatorClient()
    try:
        with open(file_path, "rb") as image_file:
            content = image_file.read()

        image = vision.Image(content=content)
        response = client.text_detection(image=image)

        if response.error.message:
            raise Exception(f"Vision API Error: {response.error.message}")

        if response.text_annotations:
            text = response.text_annotations[0].description
        return text
    except Exception as e:
        print(f"Error extracting text from image: {e}")

    return text